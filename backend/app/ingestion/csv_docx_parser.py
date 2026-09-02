import os
import csv
import pandas as pd
import docx
from typing import Dict, Any

class CSVDocxParser:
    @staticmethod
    def parse_csv(file_path: str) -> Dict[str, Any]:
        df = pd.read_csv(file_path)
        df = df.dropna(how='all')
        headers = [str(col) for col in df.columns]
        rows = df.fillna("").values.tolist()

        return {
            "page_count": 1,
            "pages": [{"page_number": 1, "text": df.to_string(), "detected_sections": []}],
            "tables": [{
                "page_number": 1,
                "sheet_name": "CSV Data",
                "table_title": os.path.basename(file_path),
                "headers": headers,
                "rows": rows,
                "raw_csv": df.to_csv(index=False)
            }],
            "full_text": df.to_csv(index=False)
        }

    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        tables_data = []
        for t_idx, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                table_rows.append([cell.text.strip() for cell in row.cells])
            if table_rows:
                headers = table_rows[0]
                data_rows = table_rows[1:]
                tables_data.append({
                    "page_number": 1,
                    "sheet_name": None,
                    "table_title": f"Document Table {t_idx + 1}",
                    "headers": headers,
                    "rows": data_rows,
                    "raw_csv": "\n".join([",".join([f'"{c}"' for c in r]) for r in table_rows])
                })

        return {
            "page_count": 1,
            "pages": [{"page_number": 1, "text": full_text, "detected_sections": []}],
            "tables": tables_data,
            "full_text": full_text
        }
